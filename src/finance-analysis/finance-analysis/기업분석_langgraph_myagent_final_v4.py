
from typing import TypedDict, List, Optional
import datetime as dt

# Langchain/Langgraph
from langgraph.graph import StateGraph, END
from langchain.prompts import ChatPromptTemplate
import networkx as nx
import matplotlib.pyplot as plt
# OpenAI API
import os
from dotenv import load_dotenv
load_dotenv()
#from langchain_community.chat_models import ChatOpenAI
from langchain_openai import ChatOpenAI
# 도구들
import pandas as pd
import yfinance as yf # yahoo finance 웹 데이터 추출 비공식 도구
from ddgs import DDGS  # DuckDuckgo 검색도구

# 회사 이름으로 yfinance 티커 추출
DART_API_KEY = os.getenv('DART_API_KEY')
import dart_fss as dart
import pandas as pd
from rapidfuzz import fuzz
from bs4 import BeautifulSoup
import requests

#  naver 뉴스 크롤링 및 관련성 검증(reranker)
import naver_latest_news_urls as newsurl
import news_maintext_extract as newstext
import sentences_embedding_reranker as rank

# 보고서 -> 워드파일 저장
# pip install python-docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


reranker = rank.BGEreranker()

# OLLAMA LLM Or OpenAI API
OLLAMA_LLM = False

# 상태 정의 ( 그래프의 데이터 공유 활용 )
class CompanyState(TypedDict):
    question: str                     # 사용자의 원 질문
    company_hint: Optional[str]       # 사용자가 전달하는 회사/티커 힌트 ,   # str 또는 None 허용, 하지만 인자 자체는 '필수'
    ticker: Optional[str]             # 최종 확인된 티커 (예: AAPL,  NVDA )
    company_name: Optional[str]       # 회사명 ( 예: Apple Inc. )
    price_df: Optional[pd.DataFrame]
    market_price_snapshot: Optional[dict]    # 시총/밸류 등 간단 스냅샷
    news: Optional[List[dict]]           # 최근 뉴스 top N개  ,  [{}, {}...{}]형태
    notes: List[str]                     # 각 노드 진행 단계 기록 / 메모
    analysis_draft: Optional[str]        # 기업분석 초안 보고서
    final_report : Optional[str]         # 기업분석 최종 보고서

# state(딕셔너리처럼 동작하는 객체)에 "notes"라는 키가 없으면 빈 리스트 []를 만들어 넣고,
# 이미 "notes" 키가 있으면  키에 해당하는 기존 value를 그대로 반환
# append(text)로 해당 리스트 끝에 text를 추가
def add_note(state:CompanyState, text:str):
    state.setdefault("notes",[]).append(text)


def safe_get(corp, key, default=None):
    """Corp 객체에서 key를 안전하게 꺼내기 (속성 → to_dict → _info 순서)"""
    # 1) 직접 속성 시도
    try:
        return getattr(corp, key)
    except AttributeError:
        pass
    # 2) to_dict() 시도
    try:
        d = corp.to_dict()
        if isinstance(d, dict) and key in d:
            return d.get(key)
    except Exception:
        pass
    # 3) 내부 info 딕셔너리 시도
    try:
        info = getattr(corp, "_info", None)
        if isinstance(info, dict) and key in info:
            return info.get(key)
    except Exception:
        pass
    return default

# print(df)

def find_stock_codes(df: pd.DataFrame, query: str, cutoff: int = 75):
    """회사명(한/영) 유사도 검색 → 상장사만 반환"""
    cand = df.copy()
    cand = cand[cand["stock_code"].astype(str).str.len() == 6]  # 상장사만

    def score_row(row):
        name_ko = row.get("corp_name") or ""
        name_en = row.get("corp_name_eng") or ""
        s1 = fuzz.WRatio(query, name_ko) if name_ko else 0
        s2 = fuzz.WRatio(query, name_en) if name_en else 0
        return max(s1, s2)

    cand["score"] = cand.apply(score_row, axis=1)
    cand = cand[cand["score"] >= cutoff].sort_values("score", ascending=False)

    return cand[["corp_name", "corp_name_eng", "corp_code", "stock_code", "score"]][:1]


# 네이버로 시장 구분 감지 → 야후 티커 접미사 생성
def detect_market_suffix_naver(stock_code: str) -> str:
    """
    네이버 종목 메인 페이지 텍스트에서 KOSPI/KOSDAQ 감지
    반환: '.KS' / '.KQ' / 또는 ''(감지 실패)
    """
    url = f"https://finance.naver.com/item/main.naver?code={stock_code}"
    r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
    # print('r : ', r)  # <Response [200]> 이면 정상
    if r.status_code != 200:
        return ""
    soup = BeautifulSoup(r.text, "lxml")
    # class="kospi" 인 img 태그 찾기
    img_tag = soup.find("img", class_=["kospi", "kosdaq"])

    if img_tag:
        alt_text = img_tag.get("alt")
        # print("alt 속성:", alt_text)
    else:
        print("해당 태그를 찾을 수 없습니다.")

    if "코스닥" in alt_text:
        return ".KQ"
    if "코스피" in alt_text:
        return ".KS"
    return ""



# 분석할 회사(티커등등) 결정
def Decision_Company(state: CompanyState):
    # hint변수에 "company_hint" 값이 있으면 그것을 저장하고 없으면 "question" 값을 저장
    com_hint = state.get("company_hint") or state["question"]
    # node 진행 단계 기록/메모
    add_note(state, f"[Decison Company Node]에서 {com_hint} 회사 추출 시도!!")

    # dart공시사이트 + 네이버금융 활용해 회사 이름으로 yfinance 티커 추출
    dart.set_api_key(api_key=DART_API_KEY)
    corp_list = dart.get_corp_list()  # 기업리스트 CorpList 객체 반환

    # Corp 객체 → 레코드
    rows = []
    for corp in corp_list:
        rows.append({
            "corp_name": safe_get(corp, "corp_name"),
            "corp_code": safe_get(corp, "corp_code"),
            "stock_code": safe_get(corp, "stock_code"),  # 상장사는 6자리, 비상장은 None/빈값
            "corp_name_eng": safe_get(corp, "corp_name_eng", None),  # 없으면 None
        })

    df = pd.DataFrame(rows)

    print(f"\n[검색어] {com_hint}")
    results = find_stock_codes(df, com_hint)  # Dataframe 반환

    if results.empty:
        print(" 국내 상장 stock code을 찾지 못했습니다.")
    else:
        results = results.to_dict(orient='records')[0]  # 데이터프레임을 “레코드(행)” 단위의 딕셔너리 리스트로 변환
        # print(results['stock_code'])

        stock_code_suffix = detect_market_suffix_naver(results['stock_code'])
        # print(stock_code_suffix)
        ticker_name = results['stock_code'] + stock_code_suffix
        print("국내 상장 종목 심볼(티커) : ", ticker_name)

        # yfinance를 이용해 티커 활용 market 정보 획득
        if ticker_name:
            info = yf.Ticker(ticker_name).get_info() or {}
            print('info : ', info) # info :  {'trailingPegRatio': None}  경우는 yfinance에서 요구하는 티커가 아닐 경우
            if info and "shortName" in info:
                print('티커 매핑 성공')
                state["ticker"] = ticker_name
                state["company_name"] = info.get("shortName")
                add_note(state, f"[Decison Company Node] 티커 확정: {ticker_name} / {state['company_name']}")
                # print('state : ', state)
                return state

    # 실패 시: 사용자 힌트 기반 그대로 진행(티커 없음)
    add_note(state, "[Decison Company Node] 티커 매핑 실패—회사명 기반 진행")
    state["company_name"] = com_hint
    # print('state : ', state)
    return state

# 회사 시장가격 수집 및 획득( 스냅샷 ) / 밸류에이션 체크 소재
def Get_MarketPrice(state: CompanyState):
    ticker = state.get("ticker")
    if not ticker: # ticker가 없는 경우
        add_note(state,"[MarketPrice] 티커 없음_가격 생략")
        return state

    add_note(state,f"[MarketPrice] {ticker} 시장 가격/정보 수집")
    pricetk = yf.Ticker(ticker)  # 시장 가격 찾을 기업 객체화
    history = pricetk.history(period="6mo", interval = "1d")
    # 수집한 histroy 가격 정보 state(상태) 객체에 저장
    state['price_df'] = history.tail(120) if not history.empty else pd.DataFrame()

    add_note(state, f"[MarketPrice] {ticker} snapshot 정보 저장")
    info = pricetk.info or {}
    market_snapshot = {
        "symbol" : ticker,
        "shortName" : info.get("shortName"),
        "longName"  : info.get("longName"),
        "currency"  : info.get("currency"), # 종목이 거래되는 통화
        "marketCap" : info.get("marketCap"), # 시가 총액
        "trailingPE": info.get("trailingPE"),  # 과거 실적 기준의 현재 주가 대비 이익 배수
        "forwardPE": info.get("forwardPE"),  # 향후 12개월 예상 EPS
        "priceToBook": info.get("priceToBook"),  # 주가 / 주당순자산(BPS).
        "fiftyTwoWeekHigh": info.get("fiftyTwoWeekHigh"),  # 52주 최고가
        "fiftyTwoWeekLow": info.get("fiftyTwoWeekLow"),  # 52주 최저가
        "sector": info.get("sector"),
        "industry": info.get("industry")
    }

    state["market_price_snapshot"] = market_snapshot
    return state


# 최근 뉴스 데이터 수집  / 최근 동향 및 뉴스 요약 소재
# ==>  DDG 서치가 아닌 네이버뉴스나 다음뉴스 최신기사를 크롤링해서
#      제공하도록 업데이트 해야함.(2025.11.05)
def Get_NewsData(state: CompanyState):
    #news_query = state.get("ticker") or state.get("company_name")
    news_query = state.get("company_hint")
    print('news_query : ', news_query)
    if not news_query:
        add_note(state, "[NewsData] 뉴스요청 없음_뉴스 생략")
        return state

    re_query = news_query + " 실적"
    add_note(state, f"[NewsData] Naver 뉴스 검색 : {re_query}")


    urllist, tavily_used = newsurl.news_latest_url(re_query)
    if tavily_used == False:
        cleaned_urls = newsurl.clean_urls(urllist)  # url이 아닌 정보나 중복 제거
    else:
        cleaned_urls = urllist
    normalized_urls = newsurl.normalize_news_urls(cleaned_urls)  # 도메인 url이 동일 정보 제거
    result = newstext.crawl_maintext_extract(normalized_urls)

    # 한글이 하나라도 없는 기사 문자열 항목 제거
    final_result = newstext.filter_valid_strings(result)
    # 최종 네이버 뉴스 본문 리스트 출력

    #print('final_result len : ', len(final_result))
    docs_split = [
        rank.MyDocument(page_content=item, metadata={"id": id}) for id, item in enumerate(final_result)
    ]

    #뉴스 크롤링 Document 리스트 항목을 벡터DB화를 위해 다시 Langchain Document 객체화
    # LangChain Document 리스트를 사용 추천 함.
    # Document ==> langchain_core.documents
    lc_docs = [rank.Document(page_content=d.page_content, metadata=d.metadata) for d in docs_split]

    print(len(lc_docs))

    # Chroma + hnsw: space = "cosine" 인 경우
    # score는 distance 임
    # 따라서 작을수록 더 유사
    vector_store = rank.lcDocument_chroma_vector_embedding(lc_docs)
    results = vector_store.similarity_search_with_score(re_query,
                                                        k=20)
    print('results Len : ', len(results))

    # 임베딩 score 로 유사 문장 Top 10 개만 추출
    embedingtop10 = rank.top_k_by_embedding_score(results, top_k=10)
    print('embedingtop10 Len : ', len(embedingtop10))

    # 10개 중 reranker로 정확성 검증해 Top5만 출력(추출)
    # print(" == rerank 결과 취합 == ")
    reranked = rank.rerank_easy(re_query, embedingtop10, reranker, top_n=5)
    print('================ reranked ===========')
    print(reranked)
    print('reranked Len : ', len(reranked))
    contextlist = []
    for i, r in enumerate(reranked):
        contextlist.append(r['doc'].page_content)

    state["news"] = contextlist


    # add_note(state,f"[NewsData] DuckDuckGo 뉴스 검색 : {news_query}")
    #
    # results = []
    # try:
    #     with DDGS() as ddgs:
    #         for r in ddgs.news(news_query,timelimit="m", max_results=10):
    #             results.append({
    #                 "title": r.get("title"),
    #                 "date": r.get("date"),
    #                 "source": r.get("source"),
    #                 "url": r.get("url"),
    #                 "body": r.get("body"),
    #             })
    # except Exception as e:
    #     add_note(state, f"[news] 검색 실패: {e}")
    #
    #state["news"] = results

    return state

# def _llm() -> ChatOpenAI:
#     # 교체 포인트 2: 모델 이름/엔드포인트만 바꿔 쓰기 쉬움
#     # temperature 0.2: 일관된 분석/요약
#     return ChatOpenAI(model="gpt-3.5-turbo", temperature=0.2)

from langchain_ollama import OllamaLLM
def _llm():
    if OLLAMA_LLM:
        llm = OllamaLLM(
            base_url='http://192.168.0.99:11434',
            model='llama3.1'  # 한국어 학습 잘된 모델
        )
    else:
        #llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.2)
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.2)
    return llm

REPORT_PROMPT = ChatPromptTemplate.from_messages([
    ("system",
     "당신은 기업 애널리스트입니다. 답변은 한국어로 간결하지만 구조적으로 작성하세요."),
    ("user",
     """아래 입력으로 기업 분석 리포트 초안을 작성하세요.

질문:
{question}

회사 식별:
- 회사명: {company_name}
- 티커: {ticker}

가격 스냅샷(JSON):
{market_price_snapshot}

6개월 가격 요약(존재 시 최근 5영업일 종가):
{recent_close_price}

최근 뉴스 Top-N(제목 - 출처 - 날짜):
{news_Top_data}

작성 지침:
1) [회사/사업 개요]를 제목으로 해서 한 단락으로 표현
2) [최근 동향]을 제목으로 해서 최근 동향 & 뉴스 요약 (불릿 3~5개)
3) [밸류에이션]을 제목으로 해서 밸류에이션 간단 체크(PE, 52주 고저 등 보수적 코멘트)
4) [리스크 요인]을 제목으로 해서 리스크 요인 3~5개
5) [한줄 결론]을 제목으로 해서 한줄 결론(투자 권유 아님, 정보 제공 목적)

배제 기준:
- 그룹 전체 전략등 그룹의 다른 계열사 관련 내용은 제외한다.
- 질문한 회사를 직접 대상으로 하지 않으면 제외한다.
- 단락 구분을 표시하는 '###'등으로 시작하는 문자열을 출력에 포함하지 마라.
"""
    )
])


# 기업분석 보고서 초안 작성
def Draft_Report(state: CompanyState):
    price_df = state.get('price_df')
    # 최근 5일 종가 처리
    if price_df is not None and not price_df.empty:
        recent_price = price_df['Close'].tail(5).round(2).to_list()
    else:
        recent_price = []

    # 뉴스 정보 처리
    news = state.get('news') or []
    # news_data = "\n".join(
    #     f"- {n.get('title','(제목없음)')} — {n.get('source')} — {n.get('date')}" for n in news[:5]
    # )
    news_data = "\n".join(news)


    llm = _llm()
    prompt = REPORT_PROMPT.format_messages(
        question = state["question"],
        company_name = state.get("company_name"),
        ticker = state.get("ticker"),
        market_price_snapshot = state.get("market_price_snapshot"),
        recent_close_price = recent_price,
        news_Top_data = news_data
    )
    # LLM에 초안 작성 요청
    draft_result = llm.invoke(prompt)

    if OLLAMA_LLM:
        state["analysis_draft"] = draft_result  # OLLAMA LLM
    else:
        state["analysis_draft"] = draft_result.content  # Openai API 동작시

    add_note(state, "[Draft Report] 초안 생성 완료")
    return state


# 기업분석 최종 보고서 작성
def Final_Report(state: CompanyState):
    draft = state.get("analysis_draft") or ""
    today = dt.date.today().isoformat()  # .isoformat() ==> date객체를 ISO 표준 문자열(YYYY-MM-DD)로 변환
    header = f"# 기업 분석 리포트\n- 기준일: {today}\n- 대상: {state.get('company_name')} ({state.get('ticker')})\n"
    state["final_report"] = f"{header}\n{draft}\n\n[유의사항]\n※ 본 자료는 투자 권유 목적이 아니며 정보 제공을 위한 요약입니다."
    add_note(state, "[final_report] 최종 보고서 생성")

    return state

def Save_Report(state: CompanyState):
    comname = state.get("company_hint")
    filepath = f"{comname}_기업분석_리포트.docx"
    report_text = state.get("final_report")
    # =========================
    # 1. Word 문서 객체 생성
    # =========================
    # python-docx의 Document 객체는 하나의 .docx 문서를 의미
    doc = Document()

    # =========================
    # 기본 스타일 설정
    # =========================
    # 'Normal' 스타일은 모든 기본 문단의 기본값
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'  # 한글 가독성을 위해 Windows 기본 한글 폰트 지정
    style.font.size = Pt(11)  # 기본 글자 크기 설정 (11pt)

    # =========================
    # 3. 문서 제목 추가
    # =========================
    # level=0 → Word에서 가장 큰 제목 스타일
    title = doc.add_heading("기업 분석 리포트", level=0)
    # 제목을 문서 중앙 정렬
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================
    # 본문 파싱 & 출력
    # =========================
    # report_text를 줄 단위로 분리하여 하나씩 처리
    for line in report_text.split("\n"):
        # 좌우 공백 제거 (불필요한 공백으로 인한 스타일 오류 방지)
        line = line.strip()

        # -------------------------
        # (1) 빈 줄 처리
        # -------------------------
        # 원문에서 줄바꿈을 유지하기 위해 빈 문단 추가
        if not line:
            doc.add_paragraph("")
            continue

        # -------------------------
        #  최상위 섹션 제목
        # -------------------------
        # 예: "1. RFHIC 기업 분석 리포트 초안"
        # line 문자열이 "1." 또는 "2." 또는 "3." 또는 "4." 또는 "5." 중 하나로 시작하면 True
        # - "기업 분석" 문구 포함 시 Heading level 1
        if line.startswith(("1.", "2.", "3.", "4.", "5.")) and "기업 분석" in line:
            doc.add_heading(line, level=1)

        # -------------------------
        # 하위 섹션 제목
        # -------------------------
        # 예: "1) 회사/사업 개요"
        # - 숫자 + ")" 형식
        elif line.startswith(("1)", "2)", "3)", "4)", "5)")):
            doc.add_heading(line, level=2)

        # 불릿 포인트
        # 예: "- RFHIC은 미국 주파수 경매 수혜 예상"
        # Word 기본 불릿(List Bullet) 스타일 적용
        elif line.startswith("-"):
            p = doc.add_paragraph(line[1:].strip())
            p.style = 'List Bullet'

        # 주의 문구
        # 예: "※ 본 자료는 투자 권유 목적이 아닙니다."
        # italic(기울임) 처리로 가독성 강화
        elif line.startswith("※"):
            p = doc.add_paragraph(line)
            p.runs[0].italic = True

        # 일반 문단
        # 위 조건에 해당하지 않는 일반 설명 문장
        else:
            doc.add_paragraph(line)

    # =========================
    # 5. Word 파일 저장
    # =========================
    # 지정한 경로에 .docx 파일로 저장
    doc.save(filepath)
    return state

# 그래프 구성
def build_graph():
    graph = StateGraph(CompanyState)
    graph.add_node("Decision_Company", Decision_Company) # 분석할 회사결정 노드 추가
    graph.add_node("Get_MarketPrice",Get_MarketPrice) # 시장가격 획득 노드 추가
    graph.add_node("Get_NewsData", Get_NewsData) # 최근 뉴스 기사 획득 노드 추가
    graph.add_node("Draft_Report", Draft_Report) # 기업분석 초안 보고서 작성 노드 추가
    graph.add_node("Final_Report", Final_Report) # 기업분석 최종 보고서 작성 노드 추가
    graph.add_node("Report_wordfile_save",Save_Report)

    graph.set_entry_point("Decision_Company")  # 그래프 시작점
    graph.add_edge("Decision_Company", "Get_MarketPrice") # 그래프 연결
    graph.add_edge("Get_MarketPrice", "Get_NewsData")  # 그래프 연결
    graph.add_edge("Get_NewsData", "Draft_Report")  # 그래프 연결
    graph.add_edge("Draft_Report", "Final_Report")  # 그래프 연결
    graph.add_edge("Final_Report", "Report_wordfile_save")  # 그래프 종점
    graph.add_edge("Report_wordfile_save", END)  # 그래프 종점

    return graph.compile()
    #return graph


# 그래프 시각화 함수 정의 ( 시각화 함수 사용시 위 그래프 컴파일을 주석 처리 해야함 )
def visualize_graph(graph: StateGraph):
    G = nx.DiGraph()

    for node in graph.nodes:
        G.add_node(node)
    for edge in graph.edges:
        G.add_edge(edge[0], edge[1])

    pos = nx.spring_layout(G)
    nx.draw(G, pos, with_labels=True, node_size=3000, node_color="lightblue", font_size=10, font_weight="bold",
            arrowsize=20)
    plt.title("StateGraph Visualization")
    plt.show()



#  main 실행
if __name__ == "__main__":
    agent_run = build_graph()

    while True:
        print()
        comname = input("투자 분석할 국내 상장 기업 이름 입력(종료 exit) : ")
        if comname == 'exit':
            break
        state1: CompanyState ={
            "question" : f"{comname} 투자 리스크 분석과 최근 뉴스 3가지만.",
            "company_hint" : f"{comname}",
            "notes" : []   # 각 노드 진행 단계 메모
        }
        output1 = agent_run.invoke(state1)
        print("="*90)
        print("\n\n", output1['final_report'])  # 기업분석 최종 보고서 상태값 출력


    # 그래프 시각화 실행
    #visualize_graph(agent_run)
